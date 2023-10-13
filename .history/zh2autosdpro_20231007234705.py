import os
import subprocess
import questionary
import rando
from docx import Document

# 定义起始路径
start_dir = "/Users/ruanliqing/Library/Mobile Documents/com~apple~CloudDocs/startup/推文/小说集"
error_dir = "/Users/ruanliqing/Library/Mobile Documents/com~apple~CloudDocs/startup/推文/处理异常的"

if not os.path.exists(error_dir):
    os.makedirs(error_dir)

#询问网址
default_webui_addr = 'http://127.0.0.1:7860'
webui_addr = input(f"Enter the webui_addr (default: {default_webui_addr}): ") or default_webui_addr

#询问是否生成视频
generate_video_input = input("Do you want to generate a video? (y/n, default: n): ").lower()
generate_video = generate_video_input in ['y', 'yes']

#询问checkpoint
def select_checkpoint():
    options = [
        "animeChangefulXL_v10ReleasedCandidate.safetensors",
        "4Guofeng4XL_v1125D.safetensors",
        "lzSDXL_v10.safetensors",
        "pixelwave_03.safetensors",
        "counterfeitxl_v10.safetensors",
        "juggernautXL_version5.safetensors",
        "sd_xl_base_1.0.safetensors"
    ]

    selected = questionary.select(
        "Please select a checkpoint:",
        choices=options
    ).ask()

    return selected

checkpoint = select_checkpoint()
print(f"Selected checkpoint: {checkpoint}")

#询问提取多少句子，做多少图
howmany = int(input(f"How many sentences do you want to process? (default: 400): ") or 400)

# 提取 英文 文件中的句子
def extract_sentences_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    sentences = '. '.join(full_text).split('. ')
    return [sentence.rstrip('.') for sentence in sentences]

# 提取 中文 文件中的句子
def extract_sentences_from_docx_zh(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # 使用中文句号、问号和感叹号来分割句子
    sentences = ' '.join(full_text)
    sentences = [s for s in re.split(r'[。？！]', sentences) if s]
    
    return sentences


# 遍历文件夹中的.docx文件
for root, dirs, files in os.walk(start_dir):
    for file in files:
        if file.endswith(".docx"):
            print(f'正在处理：{file}')

            file_path = os.path.join(root, file)

            try:
                # 从文档中提取句子
                sentences = extract_sentences_from_docx_zh(file_path)
                
                # 选择400句话
                selected_sentences = sentences if len(sentences) <= howmany else random.sample(sentences, howmany)
                
                # 写入新的.docx文件
                new_doc = Document()
                for sentence in selected_sentences:
                    new_doc.add_paragraph(sentence)
                
                # 保存新文件
                new_file_path = os.path.join("/Users/ruanliqing/Library/Mobile Documents/com~apple~CloudDocs/startup/推文/400句话英文", file.replace(".docx", "400.docx"))
                new_doc.save(new_file_path)

                new_folder_name = os.path.basename(new_file_path).replace('.docx', '')
                
                # 执行Python脚本
                script_path = "/Users/ruanliqing/Library/Mobile Documents/com~apple~CloudDocs/startup/appDeveloping/autoSD/autoSDterminal2.py"
                subprocess.run([
                    'python', script_path, 
                    '--webui_addr', webui_addr, 
                    '--new_folder_name', new_folder_name, 
                    '--docx_file_path', new_file_path,
                    '--generate_video', 'yes' if generate_video else 'no',
                    '--checkpoint', checkpoint,
                    '--start_index', '0'
                ])
                
                # 移动原文件到“已经处理完的”文件夹
                processed_dir = "/Users/ruanliqing/Library/Mobile Documents/com~apple~CloudDocs/startup/推文/已经处理完的"
                os.rename(file_path, os.path.join(processed_dir, file))

                print(f'已经处理完：{file}')

            except Exception as e:
                # 当其他异常发生时，将文件移动到“处理异常的”文件夹
                os.rename(file_path, os.path.join(error_dir, file))
                print(f'\033[91m处理异常的：{file}，原因：{e}\033[0m')
            except KeyboardInterrupt:
                # 当用户强制停止时，将文件移动到“处理异常的”文件夹
                os.rename(file_path, os.path.join(error_dir, file))
                print(f'\033[91m处理被中断的：{file}\033[0m')

        
