import requests
import os
from docx import Document

AUTH_KEY = "916ea6b7-7638-e520-0440-a4c49a4873af:fx"

def docx_zh2en(file_path):
    # 分割文件名和扩展名
    base_name, extension = os.path.splitext(os.path.basename(file_path))
    
    # 在文件名后加上"(E)"
    new_file_name = f"{base_name}(E){extension}"
    
    # 获取新的完整路径
    output_path = os.path.join(os.path.dirname(file_path), new_file_name)

    auth_key = "916ea6b7-7638-e520-0440-a4c49a4873af:fx" 
    
    # 调用zh_docx_translator函数
    en_docx_translator(file_path, output_path, "EN", auth_key)

def en_docx_translator(input_path, output_path, target_language, auth_key):
    extract_sentences_array = extract_sentences_from_docx_en(input_path)
    translated_array = array_translator(extract_sentences_array, target_language, auth_key)

    # 创建新的docx文件并添加翻译后的句子
    doc = Document()
    for sentence in translated_array:
        doc.add_paragraph(sentence)
    
    # 保存新的docx文件到指定的路径
    doc.save(output_path)

    print('翻译成功')

def zh_docx_translator(input_path, output_path, target_language, auth_key):
    extract_sentences_array = extract_sentences_from_docx_zh(input_path)
    translated_array = array_translator(extract_sentences_array, target_language, auth_key)

    # 创建新的docx文件并添加翻译后的句子
    doc = Document()
    for sentence in translated_array:
        doc.add_paragraph(sentence)
    
    # 保存新的docx文件到指定的路径
    doc.save(output_path)


def array_zh2en(sentences):
    auth_key = "916ea6b7-7638-e520-0440-a4c49a4873af:fx"
    return array_translator(sentences, "EN", auth_key)

def array_translator(sentences, target_language, auth_key):
    translated_array = []

    for sentence in sentences:
        translated_sentence = translate_with_deepl(sentence, target_language, auth_key)
        translated_array.append(translated_sentence)

    return translated_array


#测试API
def translate_with_deepl(text, target_language, auth_key):
    url = "https://api-free.deepl.com/v2/translate"
    
    payload = {
        "text": text,
        "target_lang": target_language,
        "auth_key": auth_key
    }
    
    response = requests.post(url, data=payload)
    result = response.json()
    
    translated_text = result['translations'][0]['text']
    
    return translated_text

# 提取 英文 文件中的句子，返回数组
def extract_sentences_from_docx_en(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    sentences = '. '.join(full_text).split('. ')
    return [sentence.rstrip('.') for sentence in sentences]

# 提取 中文 文件中的句子，返回数组
def extract_sentences_from_docx_zh(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # 使用中文句号、问号和感叹号来分割句子
    sentences = ' '.join(full_text)
    sentences = [s for s in re.split(r'[。？！]', sentences) if s]
    
    return sentences

if __name__ == "__main__":
    auth_key = "916ea6b7-7638-e520-0440-a4c49a4873af:fx" # 请替换为你的 API 密钥
    text_to_translate = "我爱编程"
    target_language = "EN"  # 英语

    translated_text = translate_with_deepl(text_to_translate, target_language, auth_key)
    print(translated_text)
